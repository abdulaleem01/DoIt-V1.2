# import docx NOT python-docx
import docx
from docx.shared import Inches
counter = 0
# content,image,head,doc_name


def doc_maker(*args):
    global doc_name_new
    global counter
    ###########
    if args[0] == 1:
        doc = docx.Document()
        doc.add_heading(f'{args[1]}', 0)

        doc_name_new = args[2]+".docx"
        doc.save(f'{args[2]}.docx')
        counter += 1

    if args[0] == 2:
        doc = docx.Document(doc_name_new)
        doc1 = doc.add_paragraph()
        doc1.add_run(f'{counter}.{args[1]}').bold = True
        doc.save(f'{doc_name_new}')
        counter += 1

    if args[0] == 3:
        doc = docx.Document(doc_name_new)
        doc.add_picture(f'{args[1]}', width=Inches(5.3), height=Inches(3.33))
        doc.save(f'{doc_name_new}')

    if args[0] == 4:
        doc = docx.Document(doc_name_new)
        doc1 = doc.add_paragraph()
        doc1.add_run(f'{counter}.{args[1]}').bold = True
        doc.add_picture(f'{args[2]}', width=Inches(5.3), height=Inches(3.33))
        doc.save(f'{doc_name_new}')
        counter += 1


"""

doc_maker(1,'iuhihiihuhui','fine1.docx')
doc_maker(2,'jhgjgjghghjgjjhg')
doc_maker(2,'jhgjgjghghjgjjhg')
doc_maker(2,'jhgjgjghghjgjjhg')
doc_maker(2,'jhgjgjghghjgjjhg')
doc_maker(2,'jhgjgjghghjgjjhg')
doc_maker(2,'jhgjgjghghjgjjhg')
doc_maker(3,'new.png')
doc_maker(3,'new.png')
doc_maker(3,'new.png')
doc_maker(3,'new.png')

doc_maker(2,'jhgjgjghghjgjjhg')
doc_maker(2,'jhgjgjghghjgjjhg')
doc_maker(2,'jhgjgjghghjgjjhg')
doc_maker(3,'new.png')
doc_maker(3,'new.png')

doc_maker(2,'jhgjgjghghjgjjhg')
doc_maker(2,'jhgjgjghghjgjjhg')
doc_maker(2,'jhgjgjghghjgjjhg')
doc_maker(3,'new.png')

doc_maker(2,'jhgjgjghghjgjjhg')
doc_maker(3,'new.png')

doc_maker(2,'jhgjgjghghjgjjhg')
doc_maker(2,'jhgjgjghghjgjjhg')


doc_maker(3,'new.png')
doc_maker(4,'line 2','new.png')








        
"""

###########
